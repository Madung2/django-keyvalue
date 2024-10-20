from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_COLOR_INDEX, WD_COLOR
from docx.shared import RGBColor
from ..models import Template, Rules
from .key_mapping import KeyMap
from ..serializers import RulesSerializer
import re
import io
import os
from django.conf import settings


class DocxGenerator:
    class DocxHTMLParser(HTMLParser):
        """Parse HTML TAG """
        def __init__(self, paragraph):
            super().__init__()
            self.paragraph = paragraph
            self.bold = False
            self.italic = False
            self.underline = False
            self.red =False
            self.highlight= False
            

        def handle_starttag(self, tag, attrs):
            if tag in ['b', 'strong']:
                self.bold = True
            elif tag == 'i':
                self.italic = True
            elif tag == 'u':
                self.underline = True
            elif tag== 'red':
                self.red = True
            elif tag == 'hl':
                self.highlight = True
            elif tag == 'br':
                self.paragraph.add_run().add_break()
                
        def handle_startendtag(self, tag, attrs):
            if tag == 't':
                self.paragraph.add_run().add_tab()

        def handle_endtag(self, tag):
            if tag in ['b', 'strong']:
                self.bold = False
            elif tag == 'i':
                self.italic = False
            elif tag == 'u':
                self.underline = False
            elif tag == 'red':
                self.red = False
            elif tag == 'hl':
                self.highlight = False

        def handle_data(self, data):
            #test: 
            data = data.replace('\n', '')

            if self.paragraph:
                data_splits = data.split('<t/>')
                for idx, split in enumerate(data_splits):
                    # 첫 번째 요소가 아니라면 탭 추가
                    if idx != 0:
                        self.paragraph.add_run().add_tab()
                    
                    run = self.paragraph.add_run(split)
                    if self.bold:
                        run.bold = True
                    if self.italic:
                        run.italic = True
                    if self.underline:
                        run.underline = True
                    if self.red:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    else:
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    if self.highlight:
                        run.font.highlight_color = WD_COLOR.YELLOW

###################DocxGenerator 시작##############################
    def __init__(self, data, type=3):
        for k,v in data.items():
            if v == 'N':
                data[k] = False
            elif v == 'Y':
                data[k] = True
        self.data = data
        if type == 1:
            self.contract_type = '신탁계약서'
        elif type == 2:
            self.contract_type = '전담중개업무계약서'
        elif type==3:
            self.contract_type = 'others'
        ### This is Result ###
        self.create_document()


    def _apply_html_to_paragraph(self, html_content, paragraph):
        parser = self.DocxHTMLParser(paragraph)
        parser.feed(html_content)

    def _edit_res_text(self, text):
        ##check if text inside db has ${} and replace it with target_kor_name##
        matches = re.findall(r'\$\{(.*?)\}', text)
        if matches:
            for match in matches:
                model = self.MODEL_MAP[match]
                target_kor_name = model.objects.get(id=self.data[match.lower()]).kor_name
                text = text.replace(f'${{{match}}}', target_kor_name)
            return text
        else:
            return text


    def _get_replacement_text(self, target):
        """replacement_text를 리턴한다. 조건은 2가지 번호일때, 숫자일때는 Rules 일반텍스트일때

        Args:
            target (_type_): _description_
            name (_type_): _description_
            address (_type_): _description_
            company_name (_type_): _description_

        Returns:
            _type_: _description_
        """
        test_highlight = False
        if target.isdigit():
            # 1. if Target is index number                
            rules = Rules.objects.filter(template_idx=int(target))
            rule_data = RulesSerializer(rules, many=True).data
            for rule in rule_data:
                IS_TARGET = True
                for k, v  in rule['types'].items():
                    if k.startswith('term'):                        
                        if 'other' in self.data and k.replace('term_', '') in self.data['other']:
                            IS_TARGET = True
                        else:
                            IS_TARGET = False
                            break
                    else:
                        if self.data[k] == v:
                            IS_TARGET = True
                        else:
                            IS_TARGET = False
                            break
                if IS_TARGET:
                    highlight_value = rule['highlight'] if 'highlight' in rule else False                
                    return self._edit_res_text(rule['final_text']), highlight_value

            return '', False
        else:
            # 2. if Target text do mapping

            m = KeyMap().keymap
            eng_target = m.get(target)
            if eng_target and eng_target in self.data.keys():
                return self.data[eng_target], False

            return '', False


    def create_document(self):

        # 1. generating documentlter(
        latest_template = Template.objects.all().first()
        if latest_template:
            content = latest_template.content
        else:
            raise ValueError("No template found for the given contract type.")
        
        doc = Document(content)
        ####1. 각 paragraphs 순회하면서 그 위치에 넣음
        for para in doc.paragraphs:
            targets = re.findall(r'{{\s*(.*?)\s*}}', para.text)
            for target in targets:
                #2) get replacement text
                replacement_text, highlight = self._get_replacement_text(target)
                print('target:', target, 'replace:', replacement_text )
                if replacement_text is None:
                    continue

                if replacement_text.strip() == "" and para.text.replace('{{'+target+'}}', '').strip() == "":
                    para.text = ""


                # 3) Check if the replacement text is a media file path
                if replacement_text.startswith('/media/'):
                    # Convert to absolute file path
                    image_path = os.path.join(settings.MEDIA_ROOT, replacement_text.replace('/media/', ''))
                    
                    # Check if the file exists
                    if os.path.isfile(image_path):
                        print('path is file', image_path)
                        para.clear()  # Clear the paragraph text
                        para.add_run().add_picture(image_path, width=Inches(5))  # Insert image with width of 2 inches
                        break
                    else:
                        print(f"Image not found at: {image_path}")
                        para.text = f"[Image not found: {replacement_text}]"                    

                
                if '</table>' in replacement_text: ## table이라면
                    pass
                    # table = utils.extract_table_simple(replacement_text)

                    # table_key = f'[[table{target}]]'
                    # replacement_text = replacement_text.replace(table, table_key)
                    # table_inputs[table_key] = table
                if highlight is None:
                    highlight = False
                
                # 3) replace it
                text_with_tags = para.text.replace("{{"+target+"}}", f"<hl>{replacement_text}</hl>")
                para.clear()
                
                # 4) apply html to paragraph
                self._apply_html_to_paragraph(text_with_tags, para)
    
                ##5) highlight and set font size
                for run in para.runs:
                    if highlight:
                        run.font.highlight_color = WD_COLOR.YELLOW
                    run.font.size = Pt(10)
        ## table 순회
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # for para in cell.paragraphs:  # Iterate over paragraphs within the cell
                        targets = re.findall(r'{{\s*(.*?)\s*}}', cell.text)
                        if targets ==[]:
                            continue
                        for target in targets:
                            # 2) get replacement text
                            replacement_text, highlight = self._get_replacement_text(target)
                            if replacement_text is None:
                                continue
                            if highlight is None:
                                highlight = False
                            
                            # 3) replace it
                            cell.text = cell.text.replace("{{"+target+"}}", replacement_text)                        
                        break

                                
        #########################post process start############################################

        # [[remove]] 제거 필수로 필요
        
        
        # 2. Making File itself
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        self.gen_buf = buf
        self.gen_doc = doc
