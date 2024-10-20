import re
from docx.oxml.ns import qn
from .utils import *

class NestedTableExtractor:
    def __init__(self, doc, key_value):
        self.doc = doc
        self.target_keys = [k_v['key'] for k_v in key_value if k_v['extract_all_json']]  
        # self.target_texts = fextract_data

    @staticmethod
    def remove_escape(input_text):
        # 정규 표현식을 사용하여 모든 이스케이프 문자를 찾아서 빈 문자열로 치환
        return re.sub(r'[\n\t\r\f\v]', '', input_text)
        
    @staticmethod
    def cell_has_background(cell):
        """
        Checks if the given cell has a background color.
        To make it work we need to use xml library and show namespace
        """
        cell_xml = cell._tc.get_or_add_tcPr()
        shd = cell_xml.find(qn('w:shd'))
        if shd is not None and shd.get(qn('w:fill')) != 'auto':
            return True
        return False

    def is_vertical_table(self, table):
        """
        Checks if the given table is a vertical table, meaning all cells in the first row have background color.
        """
        for cell in table.rows[0].cells:
            if not self.cell_has_background(cell):
                return False
        return True

    def extract_nested_tables(self, element):
        """
        Recursively extracts nested tables from the given document element and checks for background color in cells.
        """
        nested_tables_content = []

        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        nested_table = cell.tables[0]  # Assume only one nested table per cell
                        nested_table_data = []
                        first_row_has_color = False
                        for i, nested_row in enumerate(nested_table.rows):
                            nested_row_data = []
                            for nested_cell in nested_row.cells:
                                if nested_cell.tables:
                                    deeper_nested_tables = self.extract_nested_tables(nested_cell)
                                    nested_row_data.append(deeper_nested_tables)
                                else:
                                    if self.cell_has_background(nested_cell):
                                        nested_row_data.append((nested_cell.text, True))
                                        if i == 0:  # Check if the first row has color
                                            first_row_has_color = True
                                    else:
                                        nested_row_data.append((nested_cell.text, False))
                            nested_table_data.append(nested_row_data)
                        if first_row_has_color and self.is_vertical_table(nested_table):
                            nested_tables_content.append(nested_table_data)
                        break  # Break after finding the first nested table in the cell

        return nested_tables_content

    def extract_tables_with_background(self):
        """
        Extracts nested tables with first row colored from the given docx document.
        """
        return self.extract_nested_tables(self.doc)

    @staticmethod
    def find_target_in_tables(tables, target_text):
        """
        Finds target_text in the first column or row of each table.
        Returns the table index, location (row or column), and cell content.
        """
        results = []
        for table_index, table in enumerate(tables):
            # Check the first row
            for col_index, cell in enumerate(table[0]):
                if isinstance(cell, list):
                    continue
                cell_text, has_background = cell
                if target_text in cell_text:
                    print('TARGET:::::', (table_index, 'row', col_index, cell_text))
                    results.append((table_index, 'row', col_index, cell_text))

            # Check the first column
            for row_index, row in enumerate(table):
                if isinstance(row[0], list):
                    continue
                cell_text, has_background = row[0]
                if target_text in cell_text:
                    print('TARGET:::::', (table_index, 'column', row_index, cell_text))
                    results.append((table_index, 'column', row_index, cell_text))

        return results

    def find_table_type(self, tables, target_text):
        """
        Determines the table type based on whether '구분' is present in the target_text row or column.
        Returns 'type2' if found, otherwise 'type1'.
        """
        gu_bun_text = '구분'
        
        for table in tables:
            all_target_col = []
            gu_bun_col = -1
            
            # Find all target columns containing target_text
            for col_index, cell in enumerate(table[0]):
                if isinstance(cell, list):
                    continue
                cell_text, has_background = cell
                if target_text in cell_text:
                    all_target_col.append(col_index)
            
            # Find the gu_bun column from the target columns
            for col_index in all_target_col:
                for row in table:
                    if isinstance(row[col_index], list):
                        continue
                    cell_text, has_background = row[col_index]
                    if gu_bun_text in cell_text:
                        gu_bun_col = col_index
                        break
                if gu_bun_col != -1:
                    break
            
            # Check if gu_bun_col was found in target columns
            if gu_bun_col != -1:
                return 'type2'

        return 'type1'

    def create_dictionary_for_type1(self, table, target_text):
        """
        Creates a dictionary for type1 table where keys are below '구분' and values are below target_text.
        """
        keys = []
        values = []

        # Find the column index for '구분' and 'target_text'
        gu_bun_col = -1
        target_col = -1

        for col_index, cell in enumerate(table[0]):
            if isinstance(cell, list):
                continue
            cell_text, has_background = cell
            if cell_text.startswith('구분'):
                gu_bun_col = col_index
            if target_text in cell_text:
                target_col = col_index

        # Extract keys and values based on found column indices
        if gu_bun_col != -1 and target_col != -1:
            for row in table[1:]:
                if len(row) > gu_bun_col:
                    gu_bun_cell_text, _ = row[gu_bun_col]
                    keys.append(self.remove_escape(gu_bun_cell_text))
                if len(row) > target_col:
                    target_cell_text, _ = row[target_col]
                    values.append(self.remove_escape(target_cell_text))

        return dict(zip(keys, values))

    @staticmethod
    def split_cell_text(cell_text):
        """
        Splits the cell text by newlines and returns a list of clean texts.
        """
        return [line for line in cell_text.split('\n') if line.strip()]

    def create_dictionary_for_type2(self, table, target_text):
        """
        Creates a dictionary for type2 table where keys are in the '구분' row or column and values are in the target_text row or column.
        """
        dictionary = {}
        gu_bun_text = '구분'
        
        all_target_col = []
        target_col = -1
        gu_bun_col = -1

        # Find all target columns containing target_text
        for col_index, cell in enumerate(table[0]):
            if isinstance(cell, list):
                continue
            cell_text, has_background = cell
            if target_text in cell_text:
                all_target_col.append(col_index)

        # Find the gu_bun column from the target columns
        for col_index in all_target_col:
            for row in table:
                if isinstance(row[col_index], list):
                    continue
                cell_text, has_background = row[col_index]
                if gu_bun_text in cell_text:
                    gu_bun_col = col_index
                    break
            if gu_bun_col != -1:
                break

        # Ensure both target_col and gu_bun_col are found
        if gu_bun_col != -1:
            for row in table[1:]:
                if len(row) > gu_bun_col:
                    keys = self.split_cell_text(row[gu_bun_col][0])
                    values = self.split_cell_text(row[gu_bun_col+1][0])  # Assuming values are in the next column of gu_bun_col
                    for key, value in zip(keys, values):
                        dictionary[key] = value

        return dictionary

    def extract_all_data(self):
        """
        Extracts all data for each target key and returns a dictionary with target keys as keys and their respective data as values.
        """
        tables_with_background = self.extract_tables_with_background()
        results = {}
        
        for target_key in self.target_keys:
            target_tables = self.find_target_in_tables(tables_with_background, target_key)
            
            for table_index, _, _, _ in target_tables:
                table_type = self.find_table_type(tables_with_background, target_key)
                if table_type == 'type1':
                    results[target_key] = self.create_dictionary_for_type1(tables_with_background[table_index], target_key)
                else:
                    results[target_key] = self.create_dictionary_for_type2(tables_with_background[table_index], target_key)
        
        return results