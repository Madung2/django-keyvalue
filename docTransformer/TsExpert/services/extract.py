import re
import environ
from docx import Document
from docx.table import _Cell, Table
from fuzzywuzzy import fuzz
import tempfile
from docx.oxml.ns import qn
from .utils import *
from .find_pos import PosFinder
from .convert_pdf import convert_docx_to_pdf
from .table_detaction import DocxTableTypeDetector
from DocTransformer.settings import env


class DocxTableExtractor():
    """docx 테이블을 타입별로 분류 추출한다.
    조건1) 가로테이블
    조건2) 세로 테이블
    조건3) 가로 복수 테이블
    """
    def __init__(self, table):
        self.table = table
        self.detector = DocxTableTypeDetector(self.table)
        self.table_type_num = self.detector.type_num
        self.table_data = self.detector.table_extract
        self.extracted_key_values = {} # {k,v}
        self.extract_data()
        
    
    def _extract_horizontal_table(self):
        """가로 테이블 추출
        가로테이블은 [0]번 row가 키. 단 연달아 배경색이 있는 경우 글자 내용이 다르면 둘다 합쳐서 키로 상정한다.
        배경색이 없는 첫번째 데이터가 밸류.
        """
        print('extracted from horizontal table')
        result = {}
        
        for row_idx, row in enumerate(self.table_data):
            if 'bg'in row[0] and row[0]['bg']!=False:
                # 첫번째 열의 배경색이 있으면 키로 설정
                key = row[0]['txt']

                # 연속된 배경색이 있는 경우를 처리
                for i in range(1, len(row)):
                    if 'bg' in row[i] and row[i]['bg']!= False:
                        # 배경색이 있고, 텍스트가 다르면 키를 합친다
                        if row[i]['txt'] != key:
                            key+= " + " + row[i]['txt']
                    else:
                        # 배경색이 없는 첫번째 데이터를 값으로 설정
                        value =row[i]['txt']
                        # print('111', key, row_idx)
                        result[key] = [value, (row_idx,i)]
                        break
            else:
                # 예외적으로 한 줄에 배경색이 없는데 horizontal_table로 분류된 케이스가 있음.
                # 이 경우엔 [0]번 column이 키 [1]번 column이 value
                no_bg_found = all('bg' not in cell or cell['bg'] == False for cell in row)
                if no_bg_found and len(row) > 1:
                    key = row[0]['txt']
                    if row[1]['txt'].strip()!='':
                        value =row[1]['txt']
                        value_position = 1
                    elif len(row) > 2:
                        value = row[2]['txt']
                        value_position = 2
                    else:
                        value =''
                        value_position = None

                    result[key] = [value, (row_idx, value_position)]
                
                
        self.extracted_key_values = result

    def _extract_vertical_table(self):
        """세로 테이블 추출
        세로 테이블은 [0]번 row에 있는 항목들이 각각 키로 사용된다.
        연속된 row 같은 col에 배경색이 있다면, 해당 키들은 ' + '로 합쳐서 사용한다.
        각 col에서 배경색이 없는 첫 번째 데이터를 밸류로 설정.
        """
        print('extracted from verticle table')
        result = {}
        col_count = len(self.table_data[0])
        for col_idx in range(col_count):
            key= None
            value=None

            #같은 column에 대해 각 row를 순차적으로 탐색
            for row_idx, row in enumerate(self.table_data):
                if 'bg' in row[col_idx] and row[col_idx]['bg'] != False:
                    if key:
                        key+= ' + ' + row[col_idx]['txt']
                    else:
                        key =row[col_idx]['txt']
                elif not value and 'bg' in row[col_idx] and row[col_idx]['bg'] == False:
                    # 배경색이 없는 첫번째 데이터를 값으로 설정
                    value = row[col_idx]['txt']
                    value_position = row_idx
                    break
            # 키와 값이 모두 설정된 경우 결과에 추가
            if key and value:
                result[key] = [value, (value_position, col_idx)]
        self.extracted_key_values = result
    
    def _extract_horizontal_plural_table(self):
        """가로 복수 테이블 추출
        이 경우엔 키와 밸류가 연달아 붙어있다고 가정. 한 줄에는 복수의 키밸류 세트가 있을 수 있음.
        왼쪽에 배경색이 있는 셀이 있고 그 직후에 배경색이 없는 셀이 있으면 키밸류로 인지.  
        """
        print('extracted from horizontal plural table')
        result = {}

        for row_idx, row in enumerate(self.table_data):
            #row_result ={}

            #현재 행의 셀을 순차적으로 탐색
            col_idx = 0 
            while col_idx < len(row):
                # 현재 셀이 배경색의 가진다면, 그 셀을 키로 설정
                if 'bg' in row[col_idx] and row[col_idx]['bg'] != False:
                    key =row[col_idx]['txt']

                    # 그 다음셀의 배경색이 없으면 값을 설정
                    if col_idx +1 < len(row) and 'bg' in row[col_idx+1] and row[col_idx+1]['bg']==False:
                        value = row[col_idx+1]['txt']
                        # print('333', key, value, row_idx)
                        
                        result[key] = [value, (row_idx, col_idx+1)]
                        
                # 이번 값 처리 완료 다음 줄 넘어감
                col_idx += 1
            
            #if row_result:
            #    result.add(row_result)

        self.extracted_key_values = result
    
    def extract_data(self):

        if self.table_type_num == 1:
            self._extract_horizontal_table()
        elif self.table_type_num == 2:
            self._extract_vertical_table()
        elif self.table_type_num == 3:
            self._extract_horizontal_plural_table()
        else:
            print('------please code this unknown file type--------')


    


class KeyValueExtractor:
    def __init__(self, doc, key_value):
        """
        KeyValueExtractor를 문서 객체로 초기화합니다.
        
        :param doc: 데이터 추출을 위한 문서 객체.
        """
        self.doc = doc  # 처리할 문서
        self.data = []  # 추출된 키-값 쌍을 저장할 리스트
        self.data_keys = set()  # 발견된 고유 키를 저장할 집합
        self.rep_keys = set()  # 발견된 대표 키를 저장할 집합
        self.key_value = [k_v for k_v in key_value if not k_v['extract_all_json']] # [{"key": "차주사", "value_type": "text", "type": "string", "is_table": false, "extract_all_json": false, "synonym": {"priority": ["차주(시행사)"], "all": ["차주(시행사)"], "pattern": []}, "second_key": [], "specific": false, "sp_word": null, "value": [], "split": []}]
        self.image_key_value = [k_v for k_v in key_value if 'value_type' in k_v.keys() and k_v['value_type']== 'image']
        self.all_keys = [ele["key"] for ele in self.key_value] # ["차주사", "신탁사"]
        self.all_syn = [syn for item in self.key_value if item["synonym"] for syn in item["synonym"]["all"]]
        self.all_image_syns = [syn for item in self.image_key_value if item["synonym"] for syn in item["synonym"]["all"]]
        self.all_priority_syn = [syn for item in self.key_value if item["synonym"] for syn in item["synonym"]["priority"]]
        self.all_target_keys = {}
        for ele in self.key_value:
            if isinstance(ele.get('synonym',[]).get('all'),list):
                for syn in ele['synonym']['all']:
                    self.all_target_keys[syn] = ele['key']
                
        self.NoneTables = [item for item in self.key_value if item["is_table"] == False ]
        self.THRESHOLD = env.int('TSEXPERT_THRESHOLD',80)
    
    def filter_function(self, key):
        key = key.strip()
        all_keys= []
        for item in self.all_syn:
            if fuzz.ratio(key, item)>=self.THRESHOLD:
                rep = find_representatives(item, "all", self.key_value)
                return key, rep
        return None, None
    
    def extract_nested_table_data(self, nested_table):
        all_table_data =''
        for row in nested_table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)

            row = f"{' '.join(row_data)}: "
            all_table_data+=row
        return all_table_data

    def process_specific_value(self, r, v, k, nested_table):
        # function should check if value needs to be processed again and run process
        keyword_ele =find_keyword_ele_by_key(r, self.key_value)

        specific = keyword_ele['specific']
        sp_word = keyword_ele['sp_word'] # spword can be list
#        second_key = keyword_ele['second_key']


        if sp_word:
            if any(word in k for word in sp_word):
                RUN_TYPE=1 # dont have to find specific word 
            else:
                RUN_TYPE=2 # need to split text and find specific word

        # run specificFprocess_specific_value
        if specific:
            if RUN_TYPE ==1:
                return v, keyword_ele
            elif RUN_TYPE ==2:
                if nested_table:
                    # 여기에 해당하는 is_nested_table 을 추출해야하는 거임           
                    v = self.extract_nested_table_data(nested_table)

                # else:
                text_list = re.split(r'\s{3,}|,\s*|:\s*', v)                
                for text in text_list:
                    if any(word in text for word in sp_word):
                        # print(f'specific_running: {text},,{keyword_ele}')
                        return text, keyword_ele
        return v, keyword_ele

    def process_representatives_and_values(self, k, value, reps, data, data_keys, rep_keys, nested_table, table_number, row_num):
        print('4: 핵심 내용을 넣는 부분 process_representatives_and_values')

        for r in reps:
            v, keyword_ele = self.process_specific_value(r, value, k, nested_table)
            data.append([keyword_ele['key'], v, [table_number, row_num]])
            data_keys.add(k)
            rep_keys.add(r)

    def get_text_list(self):
        all_text_list = []
        num_paragraphs = min(25, len(self.doc.paragraphs))  # Get the minimum of 25 or the actual number of paragraphs
        for i in range(num_paragraphs):
            line_text = self.doc.paragraphs[i].text
            text_list = re.split(r'\(|\)|\/|\s{2,}', line_text)
            all_text_list += text_list
        processed_text_list = [ele for ele in all_text_list if ele != '']
        # print('processed_text_list', processed_text_list)
        return processed_text_list

    def process_none_table_keys(self):
        """
        테이블이 아닌 키-값 쌍을 처리하고 데이터 리스트에 추가합니다.
        
        :param none_tables: 키-값 쌍이 포함된 테이블이 아닌 요소들의 리스트.
        """
        none_table_data = []
        text_list = self.get_text_list()
        for ele in self.NoneTables:
            if ele["value"]!=[]: #블라인드펀드 등
                for fund_name in ele["value"]:
                    if fund_name in text_list:
                        none_table_data.append([ele['key'], fund_name, 1])
            else: #신청부팀점 신청직원
                # 1번: synonym=>priority로 찾는다
                for text in text_list:
                    if any(s in text for s in ele["synonym"]["priority"]):
                        none_table_data.append([ele['key'], text, 1])
                        break
                # pattern 으로 찾기
                pattern_list = ele['synonym']['pattern']
                if pattern_list !=[]:

                    for text in text_list:
                        if any( re.findall(p, text) for p in pattern_list):
                            # third element should be position information
                            none_table_data.append([ele['key'], text, 1])
        self.data += none_table_data  

    def find_nested_table(self, cell):
        for element in cell._element:
            if element.tag.endswith('tbl'):
                return Table(element, cell._parent._parent)
        return None
    
    def find_res_from_extracted(self,extracted_dict, table_idx):
        """
        주어진 extracted_dict에서 키를 우선순위 리스트와 비교해 대표 키를 찾고,
        해당 키와 값을 리스트에 추가하는 함수.
        
        Args:
            extracted_dict (dict): 추출된 키-값 딕셔너리 {k:[v, ()]}
            all_syn (set): 우선순위 키들이 포함된 집합
            all_target_keys (dict): 우선순위 키에 해당하는 대표 키 딕셔너리
            pos (int): 현재 위치 값

        Returns:
            list: [[대표 키, 값, 위치], [대표 키, 값, 위치], ...] 형식의 리스트
        """
        res = []
        image_res = [] 
        for k, v in extracted_dict.items():
            if k in self.all_syn: # 일반 텍스트 키밸류
                t_k = self.all_target_keys[k]
                extracted_value = v[0]
                row_pos = v[1][0]
                col_pos = v[1][1]


                res.append([t_k, extracted_value, (table_idx, row_pos, col_pos)])
        
        
        #self.image_key_value = [k_v for k_v in key_value if 'value_type' in k_v.keys() and k_v['value_type']== 'image']
            if k in self.all_image_syns:# key in self.image_key_value
                t_k = self.all_target_keys[k]
                image_res.append([t_k, "이미지", (table_idx, row_pos, col_pos)])
        
        return res, image_res

  
    def process_tables(self):
        """각 테이블 별로 돌면서 해당 키가 타겟에 있으면 키는 대표키 밸류는 그대로를 뽑는다. # 이미지랑 테이블 다 뽑음.
         returns:: [[k,v,pos], [k,v,pos]]
        """
        res = []
        image_res = []
        
        for table_idx, table in enumerate(self.doc.tables):
            ext = DocxTableExtractor(table)
            extracted_dict = ext.extracted_key_values #ext.extracted_ke_values는 키:밸류로 된 딕셔너리
            text_data, image_data=self.find_res_from_extracted(extracted_dict, table_idx)
            res+= text_data
            image_res += image_data

        return res, image_res
    


    def process_target_key_inner_table(self, res):
        """0번 셀에 타겟 텍스트가 있으면 그걸 기준으로 오른쪽에 테이블이 있는지 확인하고, 있으면 처리합니다."""
        target_text = ['대출 세부조건']  # 타겟 키워드 목록
        pos=0
        
        for table_idx, table in enumerate(self.doc.tables):
            for row in table.rows:
                # 첫 번째 셀에 타겟 텍스트가 있는지 확인
                if any(text in row.cells[0].text for text in target_text):
                    right_cell = row.cells[1]
                    
                    # 오른쪽 셀에 테이블이 있는지 확인
                    has_nested_table = False
                    for ele in right_cell._element:
                        if 'tbl' in ele.tag:
                            # 중첩된 테이블이 발견되면 이를 python-docx의 Table 객체로 변환
                            nested_table = Table(ele, self.doc)  # XML 요소를 Table 객체로 변환
                            ext = DocxTableExtractor(nested_table)  # DocxTableExtractor로 추출
                            extracted_dict = ext.extracted_key_values  # 추출된 딕셔너리 반환
                            # print('extracted', extracted_dict)  # 디버그용 출력 (필요 시 변경){k, [v, (row_idx,col_idx)]}
                            has_nested_table = True
                            inner_table_res, image_res = self.find_res_from_extracted(extracted_dict, table_idx)
                            res+= inner_table_res
                            break

                    if not has_nested_table:
                        print(f"No nested table found for target text '{row.cells[0].text}'")
        return res



    def extract_table_data_from_pdf(self):
        """
        PDF에서 테이블 데이터를 추출합니다.
        
        :return: PDF에서 추출된 테이블 데이터의 리스트.
        """
        table_data = []
        with tempfile.TemporaryDirectory() as tempdir:  # 임시 디렉토리를 생성
            tempdir = "/data"  # 임시 디렉토리 경로 설정
            temp_docx_path = f"{tempdir}/sample.docx"  # 임시 DOCX 파일 경로 설정
            temp_pdf_path = f"{tempdir}/sample.pdf"  # 임시 PDF 파일 경로 설정
            self.doc.save(temp_docx_path)  # 문서를 임시 DOCX 파일로 저장
            tables = convert_docx_to_pdf(temp_docx_path, temp_pdf_path)  # DOCX 파일을 PDF로 변환하고 테이블 추출
            if tables is not None:
                for table in tables:  # 추출된 각 테이블을 순회
                    df = table.df  # 테이블을 데이터프레임으로 변환
                    if len(df.columns) >= 2:  # 데이터프레임이 최소 2개의 열을 가지고 있는지 확인
                        extracted_data = df.iloc[:, :2].values.tolist()  # 첫 두 열의 데이터를 리스트로 추출
                        for [cell0_txt, cell1_txt] in extracted_data:  # 각 키-값 쌍을 순회
                            key = remove_unicode_characters(cell0_txt.strip())  # 첫 번째 셀에서 키를 추출
                            value = remove_unicode_characters(cell1_txt.strip())  # 두 번째 셀에서 값을 추출
                            k, reps = self.filter_function(key)  # 키를 필터링하고 대표 키를 확인
                            if k and (k not in self.data_keys):  # 키가 유효하고 이미 데이터 키 집합에 없는지 확인
                                self.process_representatives_and_values(k, value, reps, table_data, self.data_keys, self.rep_keys, nested_table=None)
                                # 대표 키와 값을 처리하고 데이터 리스트에 추가
        return table_data


    def find_value_match_first_row(self, first_row_cells, matching_row_cells, second_key_list):
        """
        이 함수는 first_row_cells에서 second_key_list에 있는 키 중 하나를 찾아서,
        해당 키와 매칭되는 matching_row_cells의 값을 반환합니다.

        매개변수:
        first_row_cells (list): 테이블 첫 번째 행의 셀 목록 (첫 번째 열 제외).
        matching_row_cells (list): 매칭된 행의 셀 목록 (첫 번째 열 제외).
        second_key_list (list): first_row_cells에서 매칭할 수 있는 키들의 리스트.

        반환값:
        str 또는 None: 매칭된 키에 해당하는 matching_row_cells의 텍스트.
                    매칭되는 키가 없으면 None을 반환합니다.

        함수 작동 방식:
        - 먼저 second_key_list에 있는 각 키에 대해 first_row_cells에서 정확히 일치하는 
        항목이 있는지 확인합니다.
        - 정확히 일치하는 키가 있으면 해당 키의 인덱스를 사용해 matching_row_cells에서 
        같은 인덱스의 값을 반환합니다.
        - 정확히 일치하는 키가 없으면, fuzz.ratio를 사용해 80% 이상의 유사성을 가진 항목을 
        찾아서 그에 해당하는 값을 반환합니다.
        """
        # Iterate over the second_key_list and find any key that matches in first_row_cells
        for key in second_key_list:
            # Try exact match first
            if key in first_row_cells:
                index = first_row_cells.index(key)
            else:
                # If no exact match, use fuzz.ratio to find approximate matches
                index = -1
                for i, cell_text in enumerate(first_row_cells):
                    if fuzz.ratio(key, cell_text) >= 70:
                        index = i
                        break  # Stop at the first match
            
            # If we found a valid index, return the corresponding matching row cell
            if index != -1 and index < len(matching_row_cells):  # Ensure index is valid and within bounds
                return matching_row_cells[index]
        
        return None  # Return None if no key from second_key_list is found


    def find_matching_rows(self, has_second_dict_elements):
        """2차원 리스트를 찾는 함수 (priority: 첫번째 키, second_key = 두번째 키)
        이 함수는 DOCX 문서의 테이블을 순회하면서 첫 번째 셀에 우선순위 키워드(priority)가 
        있는 행을 찾아, 매칭된 행에서 관련 텍스트를 추출합니다.

        매개변수:
        docx_path (str): DOCX 파일 경로.
        has_second_dict_elements (list): 우선순위 키워드('priority')와 두 번째 키('second_key')를 포함한 
                                        딕셔너리 목록. 'priority'는 행의 첫 번째 셀에서 검색할 키워드 리스트이고,
                                        'second_key'는 각 테이블의 첫 번째 행과 매칭할 키 리스트입니다.

        반환값:
        list: (first_row_cells, matching_row_cells)의 튜플 리스트를 반환합니다. 
            first_row_cells는 테이블 첫 번째 행의 셀들이고, matching_row_cells는 우선순위 키워드와 매칭된 행의 셀들입니다.

        함수 작동 방식:
        - DOCX 파일을 불러오고, 문서 내 모든 테이블을 순회합니다.
        - 각 테이블의 첫 번째 행을 가져오고, 이후 행의 첫 번째 셀과 우선순위 키워드를 비교합니다.
        - 정확한 매칭 또는 fuzz.ratio를 사용해 95% 이상의 유사도가 있는 경우 매칭된 행을 찾습니다.
        - 매칭된 행에 대해 find_value_match_first_row를 호출하여 second_key를 기준으로 값을 추출합니다.
        - 해당 값을 정리(clean_text)한 후 출력하고, 결과 리스트에 모든 행 데이터를 저장합니다.
        """
        
        result = []
        
        # Loop through all tables in the DOCX file
        for element in has_second_dict_elements:
            priority_keywords = element['synonym']['priority']
            second_keys = element['second_key']
            for table in self.doc.tables:
                if len(table.rows) == 0:
                    continue

                # Get the first row of the table
                first_row_cells = [cell.text.strip() for cell in table.rows[0].cells] [1:]

                # Loop through the rest of the rows
                for row in table.rows[1:]:
                    first_cell_text = row.cells[0].text.strip()
                    
                    # Check if the first cell text matches any priority keyword
                    if (first_cell_text in priority_keywords) or any(fuzz.ratio(first_cell_text, priority_key) >= 95 for priority_key in priority_keywords):
                        # Collect other cells in the matching row
                        matching_row_cells = [cell.text.strip() for cell in row.cells[1:]]
                        
                        # Append the result as a tuple (first row cells, matching row cells)
                        
                        #first_row = ['제조사명\n(수입업자명)', '모델명', 'KC 인증번호']
                        #matching_row = ['㈜익스프레스리프트', 'EXPRESS-\nAS380(MRL)', 'AAB10-H005-19002']

                        found_text= self.find_value_match_first_row(first_row_cells, matching_row_cells, second_keys)
                        cleaned_found_text = clean_text(found_text)
                        # found_text should be value of the target keyword
                        result.append([element['key'], cleaned_found_text, 0])


        # print('')
        return result

    def treat_second_key(self):
        """
        """
        has_secondkey_dict_elements = []
        for ele in self.key_value:
            if ele['second_key']:
                has_secondkey_dict_elements.append(ele)


        # print('%%%%%%%%%%%%%%%%%%')
        # print(has_secondkey_dict_elements)
        # print('%%%%%%%%%%%%%%%%%%')

        return self.find_matching_rows(has_secondkey_dict_elements)


    def update_table_data(self, table_data, second_key_data):
        # 각 second_key_data의 0번째 요소를 기준으로 table_data와 비교
        for sec_row in second_key_data:
            sec_key = sec_row[0]
            
            # table_data에서 0번째 요소가 겹치는 항목을 찾고 교체
            for idx, table_row in enumerate(table_data):
                if table_row[0] == sec_key:
                    table_data[idx] = sec_row  # 해당 row를 second_key_data의 row로 교체
                    break
        
        return table_data

    def extract_data(self):
        """
        문서에서 데이터를 추출합니다.
        
        :return: 추출된 모든 데이터의 리스트.[['차주사', '더블에스와이제일차', (table_pos, row_pos, col_pos)], ... ]
        """
        ## ✯ MAIN FUNCTION ✯ 
        ##1) 먼저 테이블내 키밸류를 추출한다
        table_data, image_data = self.process_tables()  # 테이블 데이터 처리
        # self.process_table_pos() #밸류가 이미지나 표xml이면 바로 추출이 어려우니  위치값을 

        # 0번 셀을 기준으로 오른쪽 셀에 테이블이 있으면 그 안에서도 process_table을 실행
        table_data = self.process_target_key_inner_table(table_data)
        ##2) 찾지 못한 값은 그냥 빈칸을 넣어준다
        table_data = self.add_not_extracted_column(table_data)
        #3) 2차원리스트 일때 값을 찾는다. 2차원 리스트를 찾는 함수 (priority: 첫번째 키, second_key = 두번째 키) 
        second_key_data= self.treat_second_key()
        table_data = self.update_table_data(table_data, second_key_data)


        ## 5) 테이블이 아닌 텍스트에서 NER로 키값 추출
        self.process_none_table_keys()  # 테이블이 아닌 키-값 쌍 처리
        self.data += table_data  # 추출된 테이블 데이터를 데이터 리스트에 추가
        print('extract_data!!!!!!', self.data)
        return self.data, image_data # 최종 데이터를 반환
    
    def add_not_extracted_column(self, table_data):
        # print('table_data', table_data)
        table_data = [ele for ele in table_data if len(ele)==3 ] 
        all_extracted_keys = [k for [k, v, pos] in table_data]
        for a in self.all_keys: 
            if a not in all_extracted_keys:
                table_data.append([a, '', 0])
        return table_data
    
    @staticmethod
    def remove_duplication(data):

        seen = set()  # 중복을 체크하기 위한 집합
        unique_data = []  # 중복이 제거된 요소를 저장할 리스트
        for ele in data:
            # ele[0]이 리스트인 경우 튜플로 변환해서 처리
            key = tuple(ele[0]) if isinstance(ele[0], list) else ele[0]
            if key not in seen:
                seen.add(key)  # 집합에 요소 추가
                unique_data.append(ele)  # 결과 리스트에도 요소 추가
        
        return unique_data
