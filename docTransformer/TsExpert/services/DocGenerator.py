from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_COLOR_INDEX
from docx.shared import RGBColor
from TsExpert.models import Template, Rules
from rest_framework.response import Response
import re
import io

keyMapping = {
    '시행사': 'developer',
    '시공사': 'constructor',
    '신탁사': 'trustee',
    '당사여신금액': 'loan_amount',
    '대출기간': 'loan_period',
    '수수료': 'fee',
    'IRR': 'irr',
    '중도상환수수료': 'prepayment_fee',
    '연채이자율': 'overdue_interest_rate',
    '원금상환유형': 'principal_repayment_type',
    '이자상환기한': 'interest_payment_period',
    '상환후불여부': 'deferred_payment',
    '연대보증금액': 'joint_guarantee_amount',
    '금융주간사': 'lead_arranger',
    '회사': 'company'
};
class DocxHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.doc = Document()
        self.paragraph = None
        self.bold = False
        self.italic = False
        self.underline = False

    def handle_starttag(self, tag, attrs):
        if tag == 'p':
            self.paragraph = self.doc.add_paragraph()
        elif tag in ['b', 'strong']:
            self.bold = True
        elif tag == 'i':
            self.italic = True
        elif tag == 'u':
            self.underline = True


    def handle_endtag(self, tag):
        if tag in ['b', 'strong']:
            self.bold = False
        elif tag == 'i':
            self.italic = False
        elif tag == 'u':
            self.underline = False

    def handle_data(self, data):
        if self.paragraph:
            # HTML에서 \t 문자를 감지하고 .docx에서 탭 문자로 변환
            data_splits = data.split('<t/>')
            for idx, split in enumerate(data_splits):
                run = self.paragraph.add_run(split)
                if self.bold:
                    run.bold = True
                if self.italic:
                    run.italic = True
                if self.underline:
                    run.underline = True

                # 마지막 요소가 아니라면 탭 추가
                if idx != len(data_splits) - 1:
                    self.paragraph.add_run().add_tab()

    def save(self, path):
        self.doc.save(path)

class DocxGenerator:

    MODEL_MAP = {
        'INVESTOR': 'InvestorModel',
        'REDEMPTION': 'RedemptionModel',
        'SUBSCRIPTION': 'SubscriptionModel',
        'TYPE': 'TypeModel',
        'BUSINESSDAY': 'BusinessDayModel',
        'MINIMUM': 'MinimumAmountModel'
    }
#     from docx.shared import RGBColor
# from docx.enum.text import WD_COLOR_INDEX
# from html.parser import HTMLParser

    # class DocxHTMLParser(HTMLParser):
    #     def __init__(self, paragraph):
    #         super().__init__()
    #         self.paragraph = paragraph
    #         self.current_run = None

    #     def add_new_run(self, text=''):
    #         self.current_run = self.paragraph.add_run(text)

    #     def handle_starttag(self, tag, attrs):
    #         self.add_new_run()  # 새로운 run 생성
    #         if tag in ['b', 'strong']:
    #             self.current_run.bold = True
    #         elif tag == 'i':
    #             self.current_run.italic = True
    #         elif tag == 'u':
    #             self.current_run.underline = True
    #         elif tag == 'red':
    #             self.current_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    #         elif tag == 'hl':
    #             self.current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    #         elif tag == 'br':
    #             self.current_run.add_break()

    #     def handle_startendtag(self, tag, attrs):
    #         if tag == 't':
    #             self.add_new_run('\t')  # 탭을 직접적으로 추가하는 대신 텍스트로 추가

    #     def handle_data(self, data):
    #         if self.current_run:
    #             self.current_run.text += data  # 현재 run에 데이터 추가

    #     def handle_endtag(self, tag):
    #         # 태그가 끝날 때 특별한 처리가 필요 없음. 모든 스타일 설정은 시작 태그에서 처리
    #         pass

    class DocxHTMLParser(HTMLParser):
        def __init__(self, paragraph):
            super().__init__()
            self.paragraph = paragraph
            self.bold = False
            self.italic = False
            self.underline = False
            self.red =False
            self.highlight= False

        def handle_starttag(self, tag, attrs):
            if tag in ['b', 'strong']:
                self.bold = True
            elif tag == 'i':
                self.italic = True
            elif tag == 'u':
                self.underline = True
            elif tag== 'red':
                self.red = True
            elif tag == 'hl':
                self.highlight = True
            elif tag == 'br':
                self.paragraph.add_run().add_break()
                
        def handle_startendtag(self, tag, attrs):
            if tag == 't':
                self.paragraph.add_run().add_tab()

        def handle_endtag(self, tag):
             
            
            if tag in ['b', 'strong']:
                self.bold = False
            elif tag == 'i':
                self.italic = False
            elif tag == 'u':
                self.underline = False
            elif tag == 'red':
                self.red = False
            elif tag == 'hl':
                self.highlight = False

        def handle_data(self, data):
            #test: 
            data = data.replace('\n', '')

            if self.paragraph:
                data_splits = data.split('<t/>')
                for idx, split in enumerate(data_splits):
                    # 첫 번째 요소가 아니라면 탭 추가
                    if idx != 0:
                        self.paragraph.add_run().add_tab()
                    
                    run = self.paragraph.add_run(split)
                    if self.bold:
                        run.bold = True
                    if self.italic:
                        run.italic = True
                    if self.underline:
                        run.underline = True
                    if self.red:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    else:
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    if self.highlight:
                        run.font.highlight_color = WD_COLOR.YELLOW



    def __init__(self, data, type):
        for k,v in data.items():
            if v == 'N':
                data[k] = False
            elif v == 'Y':
                data[k] = True
        self.data = data
        if type == 1:
            self.contract_type = '신탁계약서'
        elif type == 2:
            self.contract_type = '전담중개업무계약서'
        elif type==3:
            self.contract_type = 'others'
    
    def apply_html_to_paragraph(self, html_content, paragraph):
        print('html', html_content)
        parser = self.DocxHTMLParser(paragraph)
        parser.feed(html_content)

    def edit_res_text(self, text):
        #check if text inside db has ${} and replace it with target_kor_name
        matches = re.findall(r'\$\{(.*?)\}', text)
        if matches:
            for match in matches:
                model = self.MODEL_MAP[match]
                target_kor_name = model.objects.get(id=self.data[match.lower()]).kor_name
                text = text.replace(f'${{{match}}}', target_kor_name)
            return text
        else:
            return text

    def find_target_fields(self, rule):
        target_fields = {}
        for field in ["investor","redemption","type","businessday","minimum","subscription"]:
            field_value = getattr(rule, f"{field}_id")
            if field_value is not None:
                target_fields[field] = field_value

        for field in ["complex", "private", "foreign", "ksd", "kosdaq"]:
            if getattr(rule, field):
                target_fields[field] = True
    
        return target_fields

    def get_replacement_text(self, target, name, address, company_name):
        test_highlight = False
        if target.isdigit():
            print('is_digit')
            

            # 1. if Target is index number                
            rules = Rules.objects.filter(template_idx=int(target))
            print('rules')
            print(rules)
            
            for rule in rules:
                IS_TARGET = True
                for k, v  in rule['types'].items():
                    print('k:', k , 'v:', v )
                    ## 먼저 value를 영어화 한다
                    if type(v)!=bool and v in ALL_VALS.keys():
                        v = ALL_VALS[v]

                    if k.startswith('term'):                        
                        if 'other' in self.data and k.replace('term_', '') in self.data['other']:
                            IS_TARGET = True
                        else:
                            IS_TARGET = False
                            break
                    else:
                        if self.data[k] == v:
                            IS_TARGET = True
                        else:
                            IS_TARGET = False
                            break
                if IS_TARGET:
                    highlight_value = rule['highlight'] if 'highlight' in rule else False                
                    return self.edit_res_text(rule['final_text']), highlight_value

            return '', False
        else:
            eng_target = keyMapping[target]
            print(self.data)
            print('eng:', eng_target)
            if eng_target in self.data:
                return self.data[eng_target], False


    #     if target in ['fund_name']:
    #         return f"<b>{self.data[target]}</b>", False
    #     elif target == 'business_number':
    #         return name, False
    #     elif target == 'business_number_address':
    #         return address, False
    #     elif target in ['company_name', 'business_number_company']:
    #         return company_name, False
    #     elif target in ['date']:
    #         date_str = self.data[target]
    #         if len(date_str) == 8 and date_str.isdigit():
    #             # YYYYMMDD 형식인 경우
    #             date_obj = datetime.strptime(date_str, "%Y%m%d")
    #         elif "년" in date_str:
    # # YYYY년 MM월 DD일 형식인 경우
    #             date_obj = datetime.strptime(date_str, "%Y년 %m월 %d일")
    #         else:
    #             # YYYY-MM-DD 형식인 경우
    #             date_obj = datetime.strptime(date_str, "%Y-%m-%d")      
    #         formatted_date_str = date_obj.strftime("%Y년 %m월 %d일")
    #         return formatted_date_str, False
    #     elif target in ['date2']:
    #         date_str = self.data['date']
    #         if len(date_str) == 8 and date_str.isdigit():
    #             # YYYYMMDD 형식인 경우
    #             date_obj = datetime.strptime(date_str, "%Y%m%d")
    #         elif "년" in date_str:
    # # YYYY년 MM월 DD일 형식인 경우
    #             date_obj = datetime.strptime(date_str, "%Y년 %m월 %d일")
    #         else:
    #             # YYYY-MM-DD 형식인 경우
    #             date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    #         formatted_date_str = date_obj.strftime("%Y. %m. %d.")
    #         return f"<b>{formatted_date_str}</b>", False
    #     elif target in ['term']:
    #         if self.data['redemption'] == 'open' and self.data['term'] == '':
    #             return '신탁계약의 해지일', False
    #         # if self.data['redemption'] == 'closed':
    #         else:
    #             date_str= self.data[target]
    #             if len(date_str) == 8 and date_str.isdigit():
    #                 # YYYYMMDD 형식인 경우
    #                 date_obj = datetime.strptime(date_str, "%Y%m%d")
    #             elif "년" in date_str:
    # # YYYY년 MM월 DD일 형식인 경우
    #                 date_obj = datetime.strptime(date_str, "%Y년 %m월 %d일")
    #             else:
    #                 # YYYY-MM-DD 형식인 경우
    #                 date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    #             formatted_term_str = date_obj.strftime("%Y년 %m월 %d일")
    #             return formatted_term_str, False

    #     elif target in ['business_number']:
    #         representative = Company.objects.filter(regis_no=self.data['business_number']).first()
    #         if representative:
    #             representative = representative.representative
    #         else:
    #             representative = '홍길동'
    #         if len(representative) ==3:
                
    #             representative = '      '.join(representative)
    #         return representative, False

        
        # elif target.isdigit():
            

        #     # 1. if Target is index number                
        #     rules = Rules.objects.filter(template_idx=int(target))
        #     rule_data = RegularContractSerializer(rules, many=True).data
        #     for rule in rule_data:
        #         IS_TARGET = True
        #         for k, v  in rule['types'].items():
        #             ## 먼저 value를 영어화 한다
        #             if type(v)!=bool and v in ALL_VALS.keys():
        #                 v = ALL_VALS[v]

        #             if k.startswith('term'):                        
        #                 if 'other' in self.data and k.replace('term_', '') in self.data['other']:
        #                     IS_TARGET = True
        #                 else:
        #                     IS_TARGET = False
        #                     break
        #             else:
        #                 if self.data[k] == v:
        #                     IS_TARGET = True
        #                 else:
        #                     IS_TARGET = False
        #                     break
        #         if IS_TARGET:
        #             highlight_value = rule['highlight'] if 'highlight' in rule else False                
        #             return self.edit_res_text(rule['final_text']), highlight_value

        #     return '', False

    #     elif "." in target:
            
    #         # 4. if Target is multiple options and if you want to list then with indexes ex) {{5:11+}} => 5->template_idx, 11->indexing
    #         indexing,template_idx = target.split(".")
    #         template_idx = template_idx.strip()

    #         if indexing == '': # Normal List without numberings
    #             rules = Rules.objects.filter(template_idx=int(template_idx))
    #             rule_data = RegularContractSerializer(rules, many=True).data
    #             final_text = []
    #             for rule in rule_data:
    #                 highlight = rule['highlight'] if 'highlight' in rule else False
    #                 if highlight:
    #                     text = f"<hl>{rule['final_text']}</hl>"
    #                 else:
    #                     text = rule['final_text']
    #                 if rule['types']:
    #                     for k, v  in rule['types'].items():
    #                         if v in ALL_VALS.keys():
    #                             v = ALL_VALS[v]
    #                         if k.startswith('term'):
    #                             if'other' in self.data and k.replace('term_','') in self.data['other']:
    #                                 final_text.append(f'{self.edit_res_text(text)}')
    #                             else:
    #                                 continue
    #                         else:
    #                             if self.data[k] == v:
    #                                 final_text.append(f'{self.edit_res_text(text)}')
    #                             else:
    #                                 # Not a target
    #                                 continue
    #             res = '<br>'.join(final_text)
    #             return res, False

    #         indexing_start = int(indexing.replace(" ", "")) # List with numberings
            
    #         rules = Rules.objects.filter(template_idx=int(template_idx))
    #         rule_data = RegularContractSerializer(rules, many=True).data
    #         final_text = []
    #         for rule in rule_data:
    #             highlight = rule['highlight'] if 'highlight' in rule else False
    #             if highlight:
    #                 text = f"<hl>{rule['final_text']}</hl>"
    #             else:
    #                 text = rule['final_text']
    #             if rule['types']:
    #                 for k, v  in rule['types'].items():
    #                     if v in ALL_VALS.keys():
    #                         v = ALL_VALS[v]
    #                     if k.startswith('term'):
    #                         if'other' in self.data and k.replace('term_','') in self.data['other']:
    #                             final_text.append(f'{str(indexing_start)}. {self.edit_res_text(text)}')
    #                             indexing_start += 1
    #                         else:
    #                             continue
    #                     else:
    #                         if self.data[k] == v:
    #                             final_text.append(f'{str(indexing_start)}. {self.edit_res_text(text)}')
    #                             indexing_start += 1
    #                         else:
    #                             # Not a target
    #                             continue
    #         res = '<br>'.join(final_text)
    #         return res, False
    #     elif "/" in target:
    #         indexing, template_idx = target.split("/")
    #         indexing_start = int(indexing.replace(" ", ""))
    #         if indexing_start and indexing_start in CIRCLE.keys():
    #             rules = Rules.objects.filter(template_idx=int(template_idx)).order_by('id')
    #             rule_data = RegularContractSerializer(rules, many=True).data
    #             final_text = []
    #             for rule in rule_data:
    #                 highlight = rule['highlight'] if 'highlight' in rule else False
    #                 if highlight:
    #                     text = f"<hl>{rule['final_text']}</hl>"
    #                 else:
    #                     text = rule['final_text']
    #                 if rule['types']:
    #                     for k, v  in rule['types'].items():
    #                         if type(v)!= bool and v in ALL_VALS.keys():
    #                             v = ALL_VALS[v]
    #                         if k.startswith('term'):
    #                             if v == True:

    #                                 if 'other' in self.data and k.replace('term_','') in self.data['other']:
    #                                     final_text.append(f'{CIRCLE[indexing_start]} {self.edit_res_text(text)}')
    #                                     indexing_start += 1
    #                                 else:
    #                                     continue
    #                             if v == False:
    #                                 if 'other' in self.data and k.replace('term_','') in self.data['other']:
    #                                     continue
    #                                 else:
    #                                     final_text.append(f'{CIRCLE[indexing_start]} {self.edit_res_text(text)}')
    #                                     indexing_start += 1
    #                         else:
    #                             if self.data[k] == v:
    #                                 final_text.append(f'{CIRCLE[indexing_start]} {self.edit_res_text(text)}')
    #                                 indexing_start += 1
    #                             else:
    #                                 # Not a target
    #                                 continue
    #             res = '<br>'.join(final_text)
    #             if len(final_text) ==1:
    #                 res = res[2:]
    #             return res, False
    #         else:
    #             return '', False
    #     else: # 타겟이 일반 텍스트인 경우
    #         if target in self.data:
    #             return self.data[target], False
    #         else: 
    #             return '', False

        

    def create_document(self):
        ## number+. 
        pattern1 = re.compile(r'^(?:\s*\d{1,2}\.)')
        table_inputs = {}
        # 1. generating documentlter(
           
        latest_trust_contract_template = Template.objects.all().first()
        if latest_trust_contract_template:
            content = latest_trust_contract_template.content
        else:

            raise ValueError("No template found for the given contract type.")
        doc = Document(content)

        first_fundname_encounter = True

        # representative = Company.objects.filter(regis_no=self.data['business_number']).first()
        # print('initial data|| ', 'business_number:', self.data['business_number'], 'company_name:', self.data['company_name'] )
        # if representative:
        #     print('representative1', representative.representative)
        #     name = representative.representative
        #     address = representative.address
        #     company_name = representative.company_name
        # elif Company.objects.filter(company_name=self.data['company_name']).first():
        #     representative = Company.objects.filter(company_name=self.data['company_name']).first()
        #     print('representative2', representative.representative)
        #     name = representative.representative
        #     address = representative.address
        #     company_name = representative.company_name
        # else:
        #     print('representative3', '홍길동')
        #     name = '홍길동'
        #     address = '서울 영등포구 여의대로 108 파크원타워1'
        #     company_name = 'NH투자증권'
        # if name and len(name) ==3:
        #     name = ' '.join(name)

        for para in doc.paragraphs:
            # para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY        
            # 1) find {{ }} in paragraphs
            if '엔에이치투자증권'in para.text and "{{company_name}}" in para.text:
                para.text = para.text.replace('엔에이치투자증권', '<b>엔에이치투자증권</b>')
                para.text = para.text.replace("{{company_name}}", "<b>{{company_name}}</b>")

            targets = re.findall(r'{{\s*(.*?)\s*}}', para.text)
            for target in targets:
                print('target', target)

                #1) align to left
                if target != 'fund_name':
                    #fundname 예외 제외
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                #2) get replacement text
                name =''
                address = '' 
                company_name = '' 
                replacement_text, highlight = self.get_replacement_text(target, name, address, company_name)
                print('rep', replacement_text)
                if replacement_text.strip() == "" and para.text.replace('{{'+target+'}}', '').strip() == "":
                    para.text = "[[remove]]"

                
                if '</table>' in replacement_text: ## table이라면

                    table = utils.extract_table_simple(replacement_text)

                    table_key = f'[[table{target}]]'
                    replacement_text = replacement_text.replace(table, table_key)
                    table_inputs[table_key] = table
                if replacement_text is None:
                    replacement_text = '[RULE NOT SET]'
                if highlight is None:
                    highlight = False
                
                # 3) replace it

                text_with_tags = para.text.replace("{{"+target+"}}", replacement_text)

                para.clear()
                
                # 4) apply html to paragraph
                self.apply_html_to_paragraph(text_with_tags, para)
    
                ## remove below later
                if self.contract_type == '신탁계약서' and 'fund_name' in targets and first_fundname_encounter:
                    for run in para.runs:
                        run.font.size = Pt(14)
                        run.font.bold = True
    
                        first_fundname_encounter = False
                else:
                    for run in para.runs:
                        if highlight:
                            run.font.highlight_color = WD_COLOR.YELLOW
                        # run.font.highlight_color = WD_COLOR.YELLOW
                        run.font.size = Pt(10)
        
        #########################post process############################################

        # if self.contract_type == '신탁계약서':
        #     for para in doc.paragraphs:
        #         if '[[remove]]' in para.text:
        #             utils.delete_paragraph(para)
        # #NH투자증권 볼드
        #     # utils.bold_target_text(doc, '(시행일)', '부 칙')
        #     # doc = utils.split_one_paragraph_into_many(doc)
        #     # utils.format_numbered_paragraphs(doc,max_line_length) #=> 53으로 커팅하는것
        #     # utils.apply_distinctive_indentation(doc)
        #     # utils.apply_space_indentation(doc)
        #     # doc = utils.adjust_alignment_based_on_index(doc)
        #     # utils.align_paragraphs_to_right(doc, '집합투자업자')
        #     # utils.align_paragraphs_to_center(doc, '부 칙')
        #     # doc.save('testing_remove_empty.docx')
        #     # utils.remove_empty_paragraph(doc)
        #     # for table_key, table in table_inputs.items():
        #         # utils.add_html_content_to_docx_at_placeholder(doc, table_key, table)
        #     # utils.change_section_and_margin(doc)
        #     ##
        # else: # 전담중개업무계약서 날인 페이지 서식 적용
        #     print('upd1')
        #     utils.bold_and_underline_brackets(doc)
        #     utils.set_font_size_if_no_large_fonts(doc)
        #     utils.center_align_paragraph_with_pattern(doc, r"\d{4}\. \d{2}\. \d{2}\.")
        #     utils.apply_bold_and_font_size(doc, '회 사')

        
        
        # 2. Making File itself
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        return buf, doc


# data= {'developer': '샘플㈜ (그룹단일 : BB-, 특수금융 : A-)', 'constructor': '㈜샘플건설', 'trustee': '샘플신탁㈜', 'loan_amount': '30000000000', 'loan_period': '30', 'fee': '1000000000', 'irr': '10.96% (조달원가 6.84%, 충당금적립율 0.61% 감안 시 Net Margin 3.51%)', 'prepayment_fee': '1.0%', 'overdue_interest_rate': None, 'principal_repayment_type': '원금: 만기일시 상환', 'interest_payment_period': None, 'deferred_payment': '후불', 'joint_guarantee_amount': '130%', 'lead_arranger': '샘플은행3', 'company': None}

# gen = DocxGenerator(data, 3)
# gen.create_document()
